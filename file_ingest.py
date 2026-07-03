"""
file_ingest.py — 업로드 파일(pdf/docx/txt) → 문장 단위 텍스트 리스트.

용도:
  성취기준·지도서·기출 PDF/docx를 올리면, 본문을 추출하고
  '문장/항목 단위'로 잘라서 리스트로 돌려준다.
  이 리스트를 사용자가 검토·수정한 뒤 Record로 저장한다.

원칙:
  - 텍스트 추출은 CPU만(LLM 불필요).
  - 자동 분할은 완벽하지 않으니, 앱에서 '편집 가능한 형태'로 보여주고
    사용자가 손보게 한다(특히 성취기준은 한 줄=한 항목이 이상적).
"""
from __future__ import annotations
import io, re


# ── 파일 → 원본 텍스트 ────────────────────────────────────────
def extract_text(filename: str, raw: bytes) -> str:
    name = filename.lower()
    if name.endswith(".txt"):
        return raw.decode("utf-8", errors="ignore")
    if name.endswith(".pdf"):
        return _extract_pdf(raw)
    if name.endswith(".docx"):
        return _extract_docx(raw)
    raise RuntimeError(f"지원하지 않는 형식: {filename} (txt/pdf/docx만)")


def _extract_pdf(raw: bytes) -> str:
    import pdfplumber
    out = []
    with pdfplumber.open(io.BytesIO(raw)) as pdf:
        for page in pdf.pages:
            out.append(page.extract_text() or "")
    return "\n".join(out)


def _extract_docx(raw: bytes) -> str:
    import docx
    d = docx.Document(io.BytesIO(raw))
    parts = []
    # 문단
    for p in d.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    # 표 안의 텍스트도(성취기준이 표로 된 경우 많음)
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" ".join(cells))
    return "\n".join(parts)


# ── 원본 텍스트 → 문장/항목 단위 리스트 ───────────────────────
# 성취기준 코드 패턴: [4국05-04], [2수01-01] 등
CODE_RE = re.compile(r"\[?\d?[가-힣]{1,2}\d{2}[-–]\d{2}\]?")

def split_items(text: str, min_len: int = 6) -> list[str]:
    """
    추출 텍스트를 항목 단위로 분할.
    전략:
      1) 성취기준 코드가 보이면 코드 앞에서 끊는다(항목 경계).
      2) 그 외에는 줄바꿈 + 문장부호(. 다/음/함 등 종결) 기준으로 분할.
      3) 너무 짧은 조각(페이지번호·머리말 등)은 버린다.
    """
    text = text.replace("\r", "\n")
    # 1) 성취기준 코드 앞에 줄바꿈 삽입 → 코드 단위로 끊기 쉽게
    text = CODE_RE.sub(lambda m: "\n" + m.group(0), text)

    raw_lines = [ln.strip() for ln in text.split("\n")]
    items = []
    buf = ""
    for ln in raw_lines:
        if not ln:
            if buf:
                items.append(buf); buf = ""
            continue
        # 종결형으로 끝나면 한 항목 마감
        buf = (buf + " " + ln).strip() if buf else ln
        if re.search(r"(다|음|함|요|됨|한다|된다|이다)[.]?$", ln) or CODE_RE.match(ln):
            items.append(buf); buf = ""
    if buf:
        items.append(buf)

    # 정리: 너무 짧거나 숫자/기호뿐인 조각 제거
    cleaned = []
    for it in items:
        it = re.sub(r"\s+", " ", it).strip()
        if len(it) < min_len:
            continue
        if re.fullmatch(r"[\d\s.\-·•]+", it):   # 페이지번호 등
            continue
        cleaned.append(it)
    return cleaned


def extract_code(item: str):
    """항목 안에 성취기준 코드가 있으면 뽑아낸다(없으면 None)."""
    m = CODE_RE.search(item)
    if not m:
        return None
    code = m.group(0)
    if not code.startswith("["):
        code = "[" + code
    if not code.endswith("]"):
        code = code + "]"
    return code


def ingest(filename: str, raw: bytes):
    """
    파일 → [(문장, 감지된코드 or None), ...]
    앱에서 이 결과를 편집 테이블로 보여주고, 확정 시 Record로 저장.
    """
    text = extract_text(filename, raw)
    items = split_items(text)
    return [(it, extract_code(it)) for it in items]


# ══════════════════════════════════════════════════════════════
# 기출 시험지 → 문항 단위 분할 (초등: 통짜 시험지 대응)
# ══════════════════════════════════════════════════════════════
import re as _re

# 문항 번호 패턴들 (임용 기출에서 흔한 형식)
#  "1.", "12.", "문 3", "문3", "[3~4]", "[3∼4]", "3)", "(3)", "3번"
_Q_PATTERNS = [
    _re.compile(r"^\s*\[\s*\d+\s*[~∼-]\s*\d+\s*\]"),   # [3~4] 묶음문항
    _re.compile(r"^\s*문\s*\d+"),                       # 문 3 / 문3
    _re.compile(r"^\s*\d+\s*\."),                       # 1.
    _re.compile(r"^\s*\(\s*\d+\s*\)"),                  # (3)
    _re.compile(r"^\s*\d+\s*\)"),                       # 3)
    _re.compile(r"^\s*\d+\s*번"),                       # 3번
]

def _is_question_start(line: str) -> bool:
    return any(p.match(line) for p in _Q_PATTERNS)


def split_questions_rule(text: str, min_len: int = 10) -> list[str]:
    """
    규칙 기반 문항 분할: 문항 번호 패턴이 나오는 줄에서 새 문항 시작.
    묶음문항([3~4] 지문+여러 물음)은 하나의 덩어리로 잡힌다(사용자가 편집표에서 쪼갬).
    완벽하지 않으니 반드시 편집표 검토를 거친다.
    """
    text = text.replace("\r", "\n")
    lines = [ln.rstrip() for ln in text.split("\n")]
    questions = []
    buf = []
    for ln in lines:
        if _is_question_start(ln) and buf:
            # 이전 문항 마감
            chunk = " ".join(x.strip() for x in buf if x.strip())
            if len(chunk) >= min_len:
                questions.append(chunk)
            buf = [ln]
        else:
            buf.append(ln)
    if buf:
        chunk = " ".join(x.strip() for x in buf if x.strip())
        if len(chunk) >= min_len:
            questions.append(chunk)
    return questions


def split_questions_llm(text: str, api_key: str, model: str = "gpt-4o-mini") -> list[str]:
    """
    LLM 기반 문항 분할(옵션, 정확도↑). 묶음문항·지문+발문을 잘 묶는다.
    긴 시험지는 비용이 들 수 있음.
    """
    import json
    from openai import OpenAI
    client = OpenAI(api_key=api_key)
    prompt = (
        "다음은 임용 시험지에서 추출한 텍스트다. 이것을 '문항 단위'로 나눠라.\n"
        "각 문항은 지문·발문·조건·배점을 하나로 묶는다. 묶음문항([3~4] 등)은 각 물음으로 분리.\n"
        "머리말·쪽번호·안내문은 제외. 원문 내용을 바꾸지 말고 그대로 묶기만 하라.\n\n"
        f"[시험지 텍스트]\n{text[:8000]}\n\n"
        "출력(JSON): {\"questions\": [\"문항1 전체\", \"문항2 전체\", ...]}"
    )
    resp = client.chat.completions.create(
        model=model, messages=[{"role": "user", "content": prompt}],
        temperature=0.0,
    )
    txt = resp.choices[0].message.content
    try:
        d = json.loads(txt)
        return [q for q in d.get("questions", []) if q.strip()]
    except Exception:
        # 실패 시 규칙 기반으로 폴백
        return split_questions_rule(text)


def is_scanned_pdf(text: str) -> bool:
    """추출 텍스트가 거의 비어있으면 스캔본(OCR 필요)으로 판단."""
    return len(text.strip()) < 20
